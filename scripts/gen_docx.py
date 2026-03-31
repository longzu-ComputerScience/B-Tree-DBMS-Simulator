"""Generate User Guide/User Guide.docx — a professional Vietnamese Word document."""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

ROOT = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
OUT_DIR = os.path.join(ROOT, "User Guide")
OUT_FILE = os.path.join(OUT_DIR, "User Guide.docx")

os.makedirs(OUT_DIR, exist_ok=True)

doc = Document()

# ── Style setup ──
style = doc.styles["Normal"]
font = style.font
font.name = "Calibri"
font.size = Pt(11)
font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)

for level in range(1, 5):
    hs = doc.styles[f"Heading {level}"]
    hf = hs.font
    hf.name = "Calibri"
    hf.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)
    if level == 1:
        hf.size = Pt(20)
        hf.bold = True
    elif level == 2:
        hf.size = Pt(16)
        hf.bold = True
    elif level == 3:
        hf.size = Pt(13)
        hf.bold = True


def add_note(text, color=RGBColor(0x15, 0x60, 0x82)):
    """Add a styled note/tip block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(f"💡 {text}")
    run.font.size = Pt(10)
    run.font.italic = True
    run.font.color.rgb = color


def add_table_row(table, cells_text, bold=False):
    row = table.add_row()
    for i, txt in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(txt))
        run.font.size = Pt(10)
        run.font.name = "Calibri"
        if bold:
            run.font.bold = True


def add_code_block(text):
    """Add a monospace code block."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x20, 0x20, 0x20)


# ═══════════════════════════════════════════
# TITLE PAGE
# ═══════════════════════════════════════════

for _ in range(6):
    doc.add_paragraph("")

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("B-TREE DBMS SIMULATOR")
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x5E)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Hướng Dẫn Sử Dụng")
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x4B, 0x4B, 0x9E)

doc.add_paragraph("")

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Mô phỏng Hệ Quản Trị Cơ Sở Dữ Liệu với chỉ mục B-Tree Bậc 3")
run.font.size = Pt(12)
run.font.italic = True
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph("")
doc.add_paragraph("")

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta.add_run("Môn: Cấu Trúc Dữ Liệu và Giải Thuật Nâng Cao\nHọc kỳ 4")
run.font.size = Pt(11)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph("")

url_p = doc.add_paragraph()
url_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = url_p.add_run("🌐 Website: https://btree-dbms.vercel.app/")
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x6E, 0xBD)

doc.add_page_break()

# ═══════════════════════════════════════════
# MỤC LỤC
# ═══════════════════════════════════════════

doc.add_heading("Mục Lục", level=1)
toc_items = [
    "1. Giới thiệu dự án",
    "2. Truy cập ứng dụng",
    "3. Mô tả tính năng",
    "4. Giao diện tổng quan",
    "5. Hướng dẫn sử dụng chi tiết",
    "   5.1 Thêm sinh viên",
    "   5.2 Xoá sinh viên",
    "   5.3 Tìm kiếm theo mã sinh viên",
    "   5.4 Tìm kiếm theo họ tên",
    "   5.5 Tải dữ liệu mẫu",
    "   5.6 Xoá toàn bộ dữ liệu",
    "   5.7 Xem lịch sử thao tác",
    "6. Hiểu về B-Tree trong ứng dụng",
    "7. Cài đặt và chạy ở máy cục bộ",
    "8. Xử lý sự cố",
    "9. Bảng thuật ngữ",
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ═══════════════════════════════════════════
# 1. GIỚI THIỆU
# ═══════════════════════════════════════════

doc.add_heading("1. Giới Thiệu Dự Án", level=1)

doc.add_paragraph(
    "B-Tree DBMS Simulator là ứng dụng web giúp trực quan hoá cách "
    "một hệ quản trị cơ sở dữ liệu (DBMS) tổ chức dữ liệu bằng chỉ mục cây B-Tree. "
    "Ứng dụng mô phỏng một bảng dữ liệu sinh viên với hai chỉ mục B-Tree bậc 3 (Order-3): "
    "một theo Mã Sinh Viên (Student ID) và một theo Họ Tên (Full Name)."
)

doc.add_paragraph(
    "Người dùng có thể thêm, xoá, tìm kiếm sinh viên và quan sát trực tiếp cách "
    "cây B-Tree thay đổi: tách nút (split), mượn nút (borrow), gộp nút (merge), thu gọn "
    "gốc (root shrink). Mọi thao tác đều được ghi lại dưới dạng lịch sử có ảnh chụp "
    "trước/sau để so sánh."
)

add_note("Ứng dụng được xây dựng cho mục đích học tập — môn CTDL&GTT nâng cao, học kỳ 4.")

# ═══════════════════════════════════════════
# 2. TRUY CẬP
# ═══════════════════════════════════════════

doc.add_heading("2. Truy Cập Ứng Dụng", level=1)

doc.add_heading("Website công khai", level=2)
p = doc.add_paragraph()
run = p.add_run("https://btree-dbms.vercel.app/")
run.font.size = Pt(12)
run.font.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x6E, 0xBD)
doc.add_paragraph(
    "Truy cập đường link trên bằng trình duyệt để sử dụng ngay, không cần cài đặt."
)
add_note(
    "Phiên bản trực tuyến chỉ bao gồm giao diện frontend. "
    "Để sử dụng đầy đủ (thêm, xoá, tìm kiếm dữ liệu), "
    "cần chạy cục bộ với cả backend (xem mục 7)."
)

doc.add_heading("Chạy cục bộ (đầy đủ tính năng)", level=2)
doc.add_paragraph(
    "Để sử dụng ứng dụng với đầy đủ backend, chạy từ thư mục gốc của repository:"
)
add_code_block("npm run dev:all")
doc.add_paragraph(
    "Lệnh này khởi động đồng thời backend (port 8000) và frontend (port 3000). "
    "Sau đó truy cập http://localhost:3000 trên trình duyệt."
)

# ═══════════════════════════════════════════
# 3. TÍNH NĂNG
# ═══════════════════════════════════════════

doc.add_heading("3. Mô Tả Tính Năng", level=1)

features = [
    ("Thêm sinh viên", "Chèn bản ghi vào bảng gốc và cả hai cây B-Tree. Cây tràn sẽ tự động tách nút."),
    ("Xoá sinh viên", "Xoá khỏi cả ba cấu trúc, cây tự cân bằng bằng mượn hoặc gộp nút."),
    ("Tìm theo mã SV", "Duyệt cây ID B-Tree, hiển thị đường đi tìm kiếm từng bước."),
    ("Tìm theo họ tên", "Duyệt cây Tên B-Tree, trả về tất cả SV trùng tên (bucket)."),
    ("Cross-highlight", "Di chuột vào dòng trong bảng → highlight key tương ứng ở cả hai cây."),
    ("Lịch sử thao tác", "Ghi nhận trước/sau mỗi thao tác, kèm danh sách sự kiện B-Tree."),
    ("Dữ liệu mẫu / Reset", "Tải 7 bản ghi mẫu hoặc xoá sạch để bắt đầu lại."),
]

table = doc.add_table(rows=1, cols=2)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Light Shading Accent 1"
hdr = table.rows[0].cells
hdr[0].text = "Tính năng"
hdr[1].text = "Mô tả"
for name, desc in features:
    add_table_row(table, [name, desc])

# ═══════════════════════════════════════════
# 4. GIAO DIỆN TỔNG QUAN
# ═══════════════════════════════════════════

doc.add_heading("4. Giao Diện Tổng Quan", level=1)

doc.add_paragraph("Giao diện chia thành các phần chính:")

areas = [
    ("Thanh tiêu đề", "Tên ứng dụng, số bản ghi, số thao tác."),
    ("Bảng điều khiển trái", "3 tab: Thêm / Xoá / Tìm kiếm. Nút Tải Mẫu và Xoá Sạch."),
    ("Kết quả tìm kiếm", "Đường đi tìm kiếm và bản ghi tìm thấy."),
    ("Ghi chú tiếng Việt", "Mẹo tương tác ở cuối cột trái."),
    ("Bảng gốc (Base Table)", "Bảng sinh viên: STT, mã SV, họ tên, giới tính."),
    ("Cây ID B-Tree", "Chỉ mục sắp theo mã sinh viên."),
    ("Cây Tên B-Tree", "Chỉ mục sắp theo họ tên (hỗ trợ bucket trùng tên)."),
    ("Lịch sử thao tác", "ADD/DELETE với trạng thái trước/sau."),
]

for name, desc in areas:
    p = doc.add_paragraph()
    run_b = p.add_run(f"• {name}: ")
    run_b.bold = True
    run_b.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

# ═══════════════════════════════════════════
# 5. HƯỚNG DẪN CHI TIẾT
# ═══════════════════════════════════════════

doc.add_heading("5. Hướng Dẫn Sử Dụng Chi Tiết", level=1)

doc.add_heading("5.1 Thêm sinh viên", level=2)
steps = [
    'Chọn tab "＋ Add" ở bảng điều khiển bên trái.',
    "Nhập Mã Sinh Viên (ví dụ: S010) — phải là mã chưa tồn tại.",
    'Nhập Họ Tên (ví dụ: "Nguyen Van A").',
    "Chọn Giới tính (Male / Female / Other).",
    'Bấm "Add Student".',
    "Dòng mới trượt vào bảng kèm hiệu ứng xanh; key mới sáng lên trong cả hai cây.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")
add_note("Nếu mã SV đã tồn tại, thao tác sẽ bị từ chối và hiện lỗi màu đỏ.")

doc.add_heading("5.2 Xoá sinh viên", level=2)
steps = [
    'Chuyển sang tab "− Delete".',
    "Nhập Mã Sinh Viên cần xoá.",
    'Bấm "Delete Student".',
    "Dòng nháy đỏ rồi biến mất; cây B-Tree cập nhật tức thì.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("5.3 Tìm kiếm theo mã sinh viên", level=2)
steps = [
    'Chuyển sang tab "⌕ Search".',
    'Ở "Search by Student ID", nhập mã SV.',
    'Bấm "Search".',
    "Đường đi tìm kiếm và bản ghi tìm thấy hiện bên dưới.",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("5.4 Tìm kiếm theo họ tên", level=2)
steps = [
    'Ở "Search by Full Name", nhập họ tên.',
    'Bấm "Search".',
    "Nếu tên tồn tại, tất cả SV trùng tên hiển thị (cơ chế bucket).",
]
for i, s in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {s}")

doc.add_heading("5.5 Tải dữ liệu mẫu", level=2)
doc.add_paragraph(
    'Bấm "🎲 Load Demo" để thêm 7 bản ghi mẫu, bao gồm trường hợp tên trùng.'
)

doc.add_heading("5.6 Xoá toàn bộ dữ liệu", level=2)
doc.add_paragraph('Bấm "↺ Reset All" để xoá sạch và trở về trạng thái ban đầu.')

doc.add_heading("5.7 Xem lịch sử thao tác", level=2)
doc.add_paragraph(
    'Cuộn xuống phần "Operation History". Bấm vào mục để xem chi tiết: '
    "sự kiện B-Tree, bảng gốc trước/sau, và cây B-Tree trước/sau."
)

# ═══════════════════════════════════════════
# 6. HIỂU VỀ B-TREE
# ═══════════════════════════════════════════

doc.add_heading("6. Hiểu Về B-Tree Trong Ứng Dụng", level=1)

doc.add_paragraph("Ứng dụng sử dụng B-Tree bậc 3 (Order-3), còn gọi là cây 2-3:")

props = [
    ("Bậc (Order)", "3"),
    ("Số key tối đa mỗi nút", "2 (= bậc − 1)"),
    ("Số con tối đa mỗi nút", "3"),
    ("Số key tối thiểu (trừ gốc)", "1"),
    ("Tất cả lá", "Cùng độ sâu"),
]

table2 = doc.add_table(rows=1, cols=2)
table2.style = "Light Shading Accent 1"
hdr = table2.rows[0].cells
hdr[0].text = "Thuộc tính"
hdr[1].text = "Giá trị"
for name, val in props:
    add_table_row(table2, [name, val])

doc.add_paragraph("")
doc.add_paragraph(
    "Chèn: key vào lá → nếu tràn → tách nút, đẩy key giữa lên cha. "
    "Lặp lại cho đến khi ổn định hoặc tạo gốc mới."
)
doc.add_paragraph(
    "Xoá: xoá key → nếu thiếu → mượn (borrow) hoặc gộp (merge). "
    "Nếu gốc trống → thu gọn (root shrink)."
)
doc.add_paragraph(
    "Chỉ mục Tên: mỗi key là họ tên, value là danh sách mã SV (bucket). "
    "Thêm SV trùng tên chỉ cập nhật bucket, không thay đổi cấu trúc cây."
)

# ═══════════════════════════════════════════
# 7. CÀI ĐẶT CỤC BỘ
# ═══════════════════════════════════════════

doc.add_heading("7. Cài Đặt và Chạy Ở Máy Cục Bộ", level=1)

doc.add_heading("Yêu cầu", level=2)
doc.add_paragraph("• Python 3.10+")
doc.add_paragraph("• Node.js 18+")
doc.add_paragraph("• npm")

doc.add_heading("Bước 1 — Clone repository", level=2)
add_code_block("git clone <url-repository>\ncd DataBase-Management-System")

doc.add_heading("Bước 2 — Cài đặt dependencies", level=2)
add_code_block("pip install -r backend/requirements.txt\ncd frontend && npm install && cd ..")

doc.add_heading("Bước 3 — Khởi động (cách nhanh nhất)", level=2)
add_code_block("npm run dev:all")
doc.add_paragraph(
    "Lệnh trên chạy đồng thời backend (port 8000) và frontend (port 3000). "
    "Mở http://localhost:3000 trên trình duyệt."
)
add_note(
    "npm run dev:all là cách khuyến nghị. "
    "Nếu muốn chạy riêng: backend dùng 'npm run backend:reload', "
    "frontend dùng 'npm run dev'."
)

doc.add_heading("Chạy kiểm thử", level=2)
add_code_block("cd backend\npython -m pytest tests/ -v")
doc.add_paragraph("47 test cases (25 B-Tree unit tests + 22 integration tests).")

# ═══════════════════════════════════════════
# 8. XỬ LÝ SỰ CỐ
# ═══════════════════════════════════════════

doc.add_heading("8. Xử Lý Sự Cố", level=1)

issues = [
    ("Trang trắng / lỗi kết nối", "Backend chưa chạy. Dùng: npm run dev:all"),
    ('Lỗi "Add failed"', "Mã SV đã tồn tại hoặc trường bắt buộc bỏ trống."),
    ("Cây chồng chéo", "Zoom out bằng scroll hoặc kéo (pan)."),
    ("Backend lỗi", "pip install -r backend/requirements.txt"),
    ("Frontend lỗi", "cd frontend && npm install"),
    ("Port bị chiếm", "Tắt tiến trình trên port 3000/8000, hoặc đổi port."),
]

table3 = doc.add_table(rows=1, cols=2)
table3.style = "Light Shading Accent 1"
hdr = table3.rows[0].cells
hdr[0].text = "Vấn đề"
hdr[1].text = "Cách khắc phục"
for issue, fix in issues:
    add_table_row(table3, [issue, fix])

# ═══════════════════════════════════════════
# 9. THUẬT NGỮ
# ═══════════════════════════════════════════

doc.add_heading("9. Bảng Thuật Ngữ", level=1)

glossary = [
    ("B-Tree", "Cây cân bằng dùng trong DBMS để tìm kiếm, chèn, xoá hiệu quả."),
    ("Order-3", "Mỗi nút tối đa 3 con, 2 key."),
    ("Split", "Tách nút tràn (>2 key) thành 2, đẩy key giữa lên cha."),
    ("Merge", "Gộp nút thiếu key với anh em qua key phân cách của cha."),
    ("Borrow", "Mượn key từ anh em thông qua cha."),
    ("Root Shrink", "Gốc trống sau merge → con duy nhất thành gốc mới."),
    ("Bucket", "Danh sách mã SV cùng tên, xử lý tên trùng trong chỉ mục."),
    ("Base Table", "Bảng dữ liệu gốc — nguồn sự thật duy nhất."),
    ("Cross-highlight", "Hover → các thành phần liên quan sáng lên."),
]

table4 = doc.add_table(rows=1, cols=2)
table4.style = "Light Shading Accent 1"
hdr = table4.rows[0].cells
hdr[0].text = "Thuật ngữ"
hdr[1].text = "Giải thích"
for term, defn in glossary:
    add_table_row(table4, [term, defn])

# ── Footer ──
doc.add_paragraph("")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("— Hết —")
run.font.size = Pt(10)
run.font.italic = True
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# ── Save ──
doc.save(OUT_FILE)
print(f"✓ Created: {OUT_FILE}")
